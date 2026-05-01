# Copyright (c) Simon Niederberger.
# Distributed under the terms of the MIT License.

import io

from dash_reportbuilder.elements import (
    CaptionElement,
    HeadingElement,
    ImageElement,
    PageBreakElement,
    ParagraphElement,
)
from dash_reportbuilder.export._base import DocxTemplate
from dash_reportbuilder.model import Report

TINY_PNG_URI = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAADElEQVR4nGP4z8AAAAMBAQDJ/pLvAAAAAElFTkSuQmCC"


def test_export_docx_produces_bytes(sample_report):
    from dash_reportbuilder.export._docx import export_docx

    result = export_docx(sample_report)
    assert isinstance(result, bytes)
    assert len(result) > 100
    # Check it's a valid ZIP (docx is a ZIP)
    assert result[:2] == b"PK"


def test_export_docx_with_template(sample_report):
    from dash_reportbuilder.export._docx import export_docx

    template = DocxTemplate(font="Arial", font_size_pt=12)
    result = export_docx(sample_report, template=template)
    assert isinstance(result, bytes)
    assert result[:2] == b"PK"


def test_export_docx_empty_report():
    from dash_reportbuilder.export._docx import export_docx

    report = Report()
    result = export_docx(report)
    assert isinstance(result, bytes)
    assert result[:2] == b"PK"


# --- Additional tests ---


def _load_docx(data: bytes):
    """Helper: load bytes as a python-docx Document."""
    from docx import Document

    return Document(io.BytesIO(data))


def test_export_docx_only_images():
    from dash_reportbuilder.export._docx import export_docx

    report = Report(title="Images Only")
    report.append(ImageElement(data_uri=TINY_PNG_URI))
    report.append(ImageElement(data_uri=TINY_PNG_URI))
    report.append(ImageElement(data_uri=TINY_PNG_URI))

    result = export_docx(report)
    doc = _load_docx(result)
    assert len(doc.inline_shapes) == 3


def test_export_docx_only_text():
    from dash_reportbuilder.export._docx import export_docx

    report = Report(title="Text Only")
    report.append(HeadingElement(text="Heading 1", level=2))
    report.append(ParagraphElement(text="Body text here."))
    report.append(CaptionElement(text="A caption."))

    result = export_docx(report)
    doc = _load_docx(result)
    assert len(doc.inline_shapes) == 0
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Heading 1" in all_text
    assert "Body text here." in all_text
    assert "A caption." in all_text


def test_export_docx_image_with_caption_meta():
    from dash_reportbuilder.export._docx import export_docx

    report = Report(title="Captioned Image")
    report.append(ImageElement(data_uri=TINY_PNG_URI, caption="Figure 1: My Chart"))

    result = export_docx(report)
    doc = _load_docx(result)
    assert len(doc.inline_shapes) == 1
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Figure 1: My Chart" in all_text


def test_export_docx_image_without_caption_meta():
    from dash_reportbuilder.export._docx import export_docx

    report = Report(title="No Caption")
    report.append(ImageElement(data_uri=TINY_PNG_URI))

    result = export_docx(report)
    doc = _load_docx(result)
    assert len(doc.inline_shapes) == 1


def test_export_docx_heading_level_from_element():
    from dash_reportbuilder.export._docx import export_docx

    report = Report(title="Headings")
    report.append(HeadingElement(text="Level 1", level=1))
    report.append(HeadingElement(text="Level 3", level=3))

    result = export_docx(report)
    doc = _load_docx(result)
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(headings) >= 2
    heading_texts = [h.text for h in headings]
    assert "Level 1" in heading_texts
    assert "Level 3" in heading_texts


def test_export_docx_page_break_does_not_crash():
    from dash_reportbuilder.export._docx import export_docx

    report = Report(title="With Breaks")
    report.append(ParagraphElement(text="Before break"))
    report.append(PageBreakElement())
    report.append(ParagraphElement(text="After break"))

    result = export_docx(report)
    assert result[:2] == b"PK"


def test_export_docx_paragraph_count(sample_report):
    """The sample_report has heading, image, paragraph, caption, page_break.
    Verify the document has at least 3 text paragraphs (heading, paragraph, caption)."""
    from dash_reportbuilder.export._docx import export_docx

    result = export_docx(sample_report)
    doc = _load_docx(result)
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    assert len(non_empty) >= 3


def test_export_docx_custom_image_width():
    from dash_reportbuilder.export._docx import export_docx

    report = Report()
    report.append(ImageElement(data_uri=TINY_PNG_URI))
    template = DocxTemplate(image_width_inches=3.0)
    result = export_docx(report, template=template)
    doc = _load_docx(result)
    from docx.shared import Inches

    assert doc.inline_shapes[0].width == Inches(3.0)


def test_export_docx_all_element_types():
    """Exercise every element type in a single document without errors."""
    from dash_reportbuilder.export._docx import export_docx

    report = Report(title="All Types")
    report.append(HeadingElement(text="H", level=2))
    report.append(ImageElement(data_uri=TINY_PNG_URI))
    report.append(ParagraphElement(text="P"))
    report.append(CaptionElement(text="C"))
    report.append(PageBreakElement())

    result = export_docx(report)
    assert isinstance(result, bytes)
    assert result[:2] == b"PK"
