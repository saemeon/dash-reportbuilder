# Copyright (c) Simon Niederberger.
# Distributed under the terms of the MIT License.

"""End-to-end tests for the public extensibility API.

Covers:

- ``register_element_type`` round-trip via FileStore reload + re-render
- Native rendering pattern (``isinstance(backend, X)`` + format-native
  escape hatch), as documented in CLAUDE.md.
"""

from __future__ import annotations

import tempfile
from dataclasses import dataclass, field
from typing import Any, ClassVar

from dash_reportbuilder import (
    DocxBackend,
    HtmlBackend,
    ImageZipBackend,
    PptxBackend,
    Report,
    TypstBackend,
    register_element_type,
)
from dash_reportbuilder.elements import _new_id
from dash_reportbuilder.protocols import ReportBackend
from dash_reportbuilder.store import FileStore

# ---------------------------------------------------------------------------
# Custom element used across these tests
# ---------------------------------------------------------------------------


@register_element_type
@dataclass
class CalloutElement:
    """A custom element that uses native escape hatches per backend.

    Demonstrates the pattern from CLAUDE.md: each backend reaches for the
    most expressive escape hatch it offers, falling back to a generic
    primitive when no native option exists.
    """

    text: str
    color: str = "#2563eb"
    id: str = field(default_factory=_new_id)

    type: ClassVar[str] = "test_callout"

    def render_into(self, backend: ReportBackend) -> None:
        if isinstance(backend, TypstBackend):
            backend.append_raw(f'#block(fill: rgb("{self.color}"))[{self.text}]')
        elif isinstance(backend, HtmlBackend):
            backend.append_raw_html(
                f'<aside style="background:{self.color};padding:0.5em">'
                f"{self.text}</aside>"
            )
        elif isinstance(backend, DocxBackend):
            backend.modify(lambda doc: doc.add_paragraph(f"[CALLOUT] {self.text}"))
        else:
            backend.add_paragraph(f"[CALLOUT] {self.text}")

    def to_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "type": self.type,
            "text": self.text,
            "color": self.color,
        }

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> CalloutElement:
        return cls(
            id=data["id"],
            text=data["text"],
            color=data.get("color", "#2563eb"),
        )


# ---------------------------------------------------------------------------
# (1) register_element_type round-trip
# ---------------------------------------------------------------------------


class TestCustomElementRoundtrip:
    """Custom elements survive serialization and re-render correctly."""

    def test_to_dict_from_dict_in_memory(self):
        item = CalloutElement(text="Note", color="#ff0000", id="x")
        d = item.to_dict()
        restored = CalloutElement.from_dict(d)
        assert restored.text == "Note"
        assert restored.color == "#ff0000"
        assert restored.id == "x"

    def test_report_roundtrip_preserves_custom_type(self):
        """Report.to_dict / from_dict reconstructs CalloutElement."""
        report = Report(title="T")
        report.append(CalloutElement(text="hello", id="c1"))
        restored = Report.from_dict(report.to_dict())
        assert len(restored.items) == 1
        assert isinstance(restored.items[0], CalloutElement)
        assert restored.items[0].text == "hello"
        assert restored.items[0].id == "c1"

    def test_filestore_reload_preserves_custom_type(self):
        """FileStore save → fresh FileStore → re-read recovers the type."""
        with tempfile.TemporaryDirectory() as tmpdir:
            store = FileStore(tmpdir)
            report = Report(title="T")
            report.append(CalloutElement(text="persisted", color="#abc", id="c1"))
            store.put("session", report)

            # Fresh store instance — proves the registry is consulted on
            # deserialization (not just held in the original object).
            store2 = FileStore(tmpdir)
            loaded = store2.get("session")
            assert len(loaded.items) == 1
            assert isinstance(loaded.items[0], CalloutElement)
            assert loaded.items[0].text == "persisted"
            assert loaded.items[0].color == "#abc"

    def test_filestore_reload_then_export_html(self):
        """After reload, the custom element still renders via render_into."""
        with tempfile.TemporaryDirectory() as tmpdir:
            store = FileStore(tmpdir)
            report = Report(title="T")
            report.append(CalloutElement(text="hello", color="#abc"))
            store.put("session", report)

            loaded = FileStore(tmpdir).get("session")
            backend = HtmlBackend(title=loaded.title)
            data = loaded.export(backend)
            html = data.decode("utf-8")
            assert "<aside" in html
            assert "background:#abc" in html
            assert "hello" in html


# ---------------------------------------------------------------------------
# (5) Native escape-hatch dispatch
# ---------------------------------------------------------------------------


class TestNativeDispatch:
    """Each backend receives output via the appropriate escape hatch."""

    def test_typst_uses_append_raw(self):
        backend = TypstBackend(title="T")
        report = Report(title="T")
        report.append(CalloutElement(text="hi", color="#123456"))
        for item in report.items:
            item.render_into(backend)
        source = backend.build_source()
        assert "#block(fill: " in source
        assert "#123456" in source
        assert "hi" in source

    def test_html_uses_append_raw_html(self):
        backend = HtmlBackend(title="T")
        report = Report(title="T")
        report.append(CalloutElement(text="hi", color="#abc"))
        for item in report.items:
            item.render_into(backend)
        source = backend.build_source()
        # Verbatim insertion — the angle brackets are NOT escaped.
        assert "<aside style=" in source
        assert "background:#abc" in source

    def test_docx_uses_modify(self):
        report = Report(title="T")
        report.append(CalloutElement(text="hi"))
        backend = DocxBackend(title=report.title)
        data = report.export(backend)
        # Round-trip the .docx to verify the paragraph landed there.
        import io

        from docx import Document

        doc = Document(io.BytesIO(data))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[CALLOUT] hi" in all_text

    def test_pptx_falls_back_to_primitive(self):
        """No PptxBackend branch → CalloutElement falls through to add_paragraph."""
        report = Report(title="T")
        report.append(CalloutElement(text="hi"))
        backend = PptxBackend()
        data = report.export(backend)
        assert data[:2] == b"PK"
        # We don't introspect the .pptx structure here — the assertion that
        # matters is "doesn't crash on an unknown element type".

    def test_image_zip_falls_back_silently(self):
        """ImageZipBackend.add_paragraph is a no-op, so the zip stays empty."""
        report = Report(title="T")
        report.append(CalloutElement(text="hi"))
        backend = ImageZipBackend()
        data = report.export(backend)
        import io
        import zipfile

        assert zipfile.ZipFile(io.BytesIO(data)).namelist() == []
