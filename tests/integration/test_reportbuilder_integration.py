# Copyright (c) Simon Niederberger.
# Distributed under the terms of the MIT License.

"""Integration tests for dash-reportbuilder — live Dash app with Chrome.

Run locally with:
  PATH="/opt/homebrew/bin:$PATH" uv run pytest .tests/integration/test_reportbuilder_integration.py -v
"""

from __future__ import annotations

import time

import pytest

import dash
import plotly.graph_objects as go
from dash import dcc, html
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

from dash_capture import capture_graph, plotly_strategy
from dash_reportbuilder import MemoryStore, report_action, report_viewer


def _make_figure():
    return go.Figure(
        data=go.Scatter(x=[1, 2, 3], y=[4, 5, 6], mode="markers"),
        layout=dict(title="Test Chart", width=400, height=300),
    )


def _find_button_in(driver, label):
    """Find a visible (interactable) button under *driver* by exact text match."""
    for b in driver.find_elements(By.TAG_NAME, "button"):
        if not b.is_displayed():
            continue
        text = driver.execute_script("return arguments[0].textContent", b)
        if text.strip() == label:
            return b
    return None


def _find_button(dash_duo, label):
    return _find_button_in(dash_duo.driver, label)


def _open_report_builder(dash_duo):
    """Open the Report Builder wizard so its inner buttons become interactable."""
    trigger = _find_button(dash_duo, "Report Builder")
    assert trigger is not None, "'Report Builder' trigger not found"
    trigger.click()
    # Wait for the modal to render (the '+ Heading' button becomes visible)
    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: _find_button_in(d, "+ Heading") is not None
    )


def _build_app():
    """Build a test app with graph + capture + report viewer."""
    app = dash.Dash(__name__)
    store = MemoryStore()

    graph = dcc.Graph(id="test-graph", figure=_make_figure())
    # Use a unique trigger label so it doesn't collide with the report
    # builder's internal "Export" button.
    wizard = capture_graph(
        graph,
        trigger="Capture",
        strategy=plotly_strategy(),
        actions=[report_action(store)],
    )

    app.layout = html.Div(
        [
            graph,
            wizard,
            html.Div(id="report-panel", children=[report_viewer(store)]),
        ]
    )
    return app, store


# ── tests ────────────────────────────────────────────────────────────────


def test_add_heading_button(dash_duo):
    """Clicking '+ Heading' adds a heading item to the viewer."""
    app, store = _build_app()
    dash_duo.start_server(app)
    dash_duo.wait_for_element("#test-graph", timeout=10)
    _open_report_builder(dash_duo)

    btn = _find_button(dash_duo, "+ Heading")
    assert btn is not None, "'+ Heading' button not found"
    btn.click()

    # Wait for a heading item to appear (drb-item class)
    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: len(d.find_elements(By.CSS_SELECTOR, ".drb-item")) >= 1
    )
    items = dash_duo.driver.find_elements(By.CSS_SELECTOR, ".drb-item")
    assert len(items) == 1

    # Verify it's in the store
    report = store.get("default")
    assert len(report.items) == 1
    assert report.items[0].type == "heading"
    assert report.items[0].text == "Section Title"


def test_add_paragraph_button(dash_duo):
    """Clicking '+ Paragraph' adds a paragraph item."""
    app, store = _build_app()
    dash_duo.start_server(app)
    dash_duo.wait_for_element("#test-graph", timeout=10)
    _open_report_builder(dash_duo)

    btn = _find_button(dash_duo, "+ Paragraph")
    assert btn is not None
    btn.click()

    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: len(d.find_elements(By.CSS_SELECTOR, ".drb-item")) >= 1
    )
    report = store.get("default")
    assert len(report.items) == 1
    assert report.items[0].type == "paragraph"


def test_add_multiple_items(dash_duo):
    """Adding multiple items shows them all in the viewer."""
    app, store = _build_app()
    dash_duo.start_server(app)
    dash_duo.wait_for_element("#test-graph", timeout=10)
    _open_report_builder(dash_duo)

    # Add heading
    _find_button(dash_duo, "+ Heading").click()
    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: len(d.find_elements(By.CSS_SELECTOR, ".drb-item")) >= 1
    )

    # Add paragraph
    _find_button(dash_duo, "+ Paragraph").click()
    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: len(d.find_elements(By.CSS_SELECTOR, ".drb-item")) >= 2
    )

    # Add caption
    _find_button(dash_duo, "+ Caption").click()
    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: len(d.find_elements(By.CSS_SELECTOR, ".drb-item")) >= 3
    )

    report = store.get("default")
    assert len(report.items) == 3
    assert [it.type for it in report.items] == ["heading", "paragraph", "caption"]


def test_delete_item(dash_duo):
    """Clicking delete removes an item from the report."""
    app, store = _build_app()
    dash_duo.start_server(app)
    dash_duo.wait_for_element("#test-graph", timeout=10)
    _open_report_builder(dash_duo)

    # Add two items
    _find_button(dash_duo, "+ Heading").click()
    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: len(d.find_elements(By.CSS_SELECTOR, ".drb-item")) >= 1
    )
    _find_button(dash_duo, "+ Paragraph").click()
    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: len(d.find_elements(By.CSS_SELECTOR, ".drb-item")) >= 2
    )
    assert len(store.get("default").items) == 2

    # Delete the first item (click the X button inside the first .drb-item)
    first_item = dash_duo.driver.find_elements(By.CSS_SELECTOR, ".drb-item")[0]
    delete_btn = first_item.find_element(By.TAG_NAME, "button")
    delete_btn.click()

    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: len(d.find_elements(By.CSS_SELECTOR, ".drb-item")) == 1
    )
    report = store.get("default")
    assert len(report.items) == 1
    assert report.items[0].type == "paragraph"


def test_reorder_persists_to_export():
    """Reorder via store.reorder() is reflected in export output.

    This tests the Python-side flow: reorder items in the store,
    then verify the export respects the new order.  The JS drag
    integration (SortableJS → set_props → order store) is tested
    separately by manual testing since dash_clientside.set_props
    cannot be reliably triggered from Selenium.
    """
    from dash_reportbuilder import HeadingElement, MemoryStore
    from dash_reportbuilder.export._docx import export_docx
    from docx import Document
    import io

    store = MemoryStore()
    report = store.get("default")
    report.append(HeadingElement(text="First", level=2, id="aaa"))
    report.append(HeadingElement(text="Second", level=2, id="bbb"))
    report.append(HeadingElement(text="Third", level=2, id="ccc"))
    store.put("default", report)

    # Reorder: reverse
    report.reorder(["ccc", "bbb", "aaa"])
    store.put("default", report)

    # Export
    data = export_docx(report)
    doc = Document(io.BytesIO(data))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    # First heading is the auto-title, then our items
    assert "Third" in headings
    assert "Second" in headings
    assert "First" in headings
    # Verify order: Third before Second before First
    idx_third = headings.index("Third")
    idx_second = headings.index("Second")
    idx_first = headings.index("First")
    assert idx_third < idx_second < idx_first


def test_sortablejs_initializes(dash_duo):
    """SortableJS is initialized on the list element."""
    app, store = _build_app()
    dash_duo.start_server(app)
    dash_duo.wait_for_element("#test-graph", timeout=10)
    _open_report_builder(dash_duo)

    # Add items so the list is non-empty
    _find_button(dash_duo, "+ Heading").click()
    _find_button(dash_duo, "+ Paragraph").click()
    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: len(d.find_elements(By.CSS_SELECTOR, ".drb-item")) >= 2
    )

    # Verify SortableJS was initialized
    has_sortable = dash_duo.driver.execute_script("""
        var list = document.querySelector('.drb-sortable-list');
        return list && list._drbSortable != null;
    """)
    assert has_sortable, "SortableJS not initialized on list element"

    # Verify order store ref exists
    ref_text = dash_duo.driver.execute_script("""
        var ref = document.querySelector('.drb-order-store-ref');
        return ref ? ref.textContent.trim() : null;
    """)
    assert ref_text is not None, "Order store ref not found"
    assert "_order" in ref_text


def test_export_docx(dash_duo):
    """Export button triggers a download (no error)."""
    app, store = _build_app()
    dash_duo.start_server(app)
    dash_duo.wait_for_element("#test-graph", timeout=10)
    _open_report_builder(dash_duo)

    # Add an item first
    _find_button(dash_duo, "+ Heading").click()
    WebDriverWait(dash_duo.driver, 10).until(
        lambda d: len(d.find_elements(By.CSS_SELECTOR, ".drb-item")) >= 1
    )

    # Click export (default format is docx)
    export_btn = _find_button(dash_duo, "Export")
    assert export_btn is not None
    export_btn.click()

    # Just verify no JS errors
    time.sleep(2)
    logs = dash_duo.driver.get_log("browser")
    errors = [l for l in logs if l["level"] == "SEVERE"]
    assert len(errors) == 0, f"Browser errors: {errors}"
