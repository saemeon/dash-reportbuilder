# Copyright (c) Simon Niederberger.
# Distributed under the terms of the MIT License.

"""Word (.docx) export via python-docx."""

from __future__ import annotations

import io

from dash_reportbuilder.export._base import DocxTemplate, decode_data_uri
from dash_reportbuilder.model import ItemType, Report


def _apply_font(run, *, font: str, size_pt: int, bold: bool = False, italic: bool = False) -> None:
    """Apply font settings to a python-docx Run."""
    from docx.shared import Pt

    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic


def export_docx(report: Report, template: DocxTemplate | None = None) -> bytes:
    """Export *report* to .docx bytes.

    When *template* provides a ``template_path``, that .docx file is opened
    as the base document — all existing styles (headers, footers, fonts,
    margins, cover page) are preserved, and report items are appended after
    the template content.

    Requires ``python-docx`` (install with ``pip install dash-reportbuilder[docx]``).
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as e:
        raise ImportError(
            "python-docx is required for Word export. "
            "Install with: pip install dash-reportbuilder[docx]"
        ) from e

    t = template or DocxTemplate()

    if t.template_path:
        doc = Document(t.template_path)
    else:
        doc = Document()
        # Apply default margins
        for section in doc.sections:
            section.top_margin = Inches(t.margin_inches)
            section.bottom_margin = Inches(t.margin_inches)
            section.left_margin = Inches(t.margin_inches)
            section.right_margin = Inches(t.margin_inches)

        # Add a report title as Heading 1
        h = doc.add_heading(report.title, level=1)
        for run in h.runs:
            run.font.name = t.font

    for item in report.items:
        if item.type == ItemType.IMAGE:
            img_bytes = decode_data_uri(item.content)
            buf = io.BytesIO(img_bytes)
            # Center the image
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(buf, width=Inches(t.image_width_inches))
            # Add caption from meta if present
            caption = item.meta.get("caption")
            if caption:
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_p.add_run(caption)
                _apply_font(cap_run, font=t.font, size_pt=t.font_size_pt - 1, italic=True)

        elif item.type == ItemType.HEADING:
            level = item.meta.get("heading_level", 2)
            h = doc.add_heading(item.content, level=level)
            for run in h.runs:
                run.font.name = t.font

        elif item.type == ItemType.PARAGRAPH:
            p = doc.add_paragraph()
            run = p.add_run(item.content)
            _apply_font(run, font=t.font, size_pt=t.font_size_pt)

        elif item.type == ItemType.CAPTION:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item.content)
            _apply_font(run, font=t.font, size_pt=t.font_size_pt - 1, italic=True)

        elif item.type == ItemType.PAGE_BREAK:
            doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
