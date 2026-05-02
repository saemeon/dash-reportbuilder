# Copyright (c) Simon Niederberger.
# Distributed under the terms of the MIT License.

"""Word (.docx) backend via python-docx."""

from __future__ import annotations

import io
from collections.abc import Callable
from typing import TYPE_CHECKING

from dash_reportbuilder.export._base import DocxTemplate, decode_data_uri

if TYPE_CHECKING:
    from docx.document import Document


class DocxBackend:
    """Builds a Word document incrementally.

    Implements the :class:`~dash_reportbuilder.protocols.ReportBackend`
    protocol.  Elements that need full python-docx control can use the
    :meth:`modify` escape hatch.

    Parameters
    ----------
    template : DocxTemplate, optional
        Template settings.  When ``template_path`` is set, that file is
        opened as the base document.
    title : str, optional
        Title for the report.  Added as a top-level heading when no
        template path is set.
    """

    def __init__(
        self,
        template: DocxTemplate | None = None,
        *,
        title: str | None = None,
    ) -> None:
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError as e:
            raise ImportError(
                "python-docx is required for Word export. "
                "Install with: pip install dash-reportbuilder[docx]"
            ) from e

        self.template = template or DocxTemplate()
        self.title = title

        if self.template.template_path:
            self.doc = Document(self.template.template_path)
        else:
            self.doc = Document()
            for section in self.doc.sections:
                section.top_margin = Inches(self.template.margin_inches)
                section.bottom_margin = Inches(self.template.margin_inches)
                section.left_margin = Inches(self.template.margin_inches)
                section.right_margin = Inches(self.template.margin_inches)
            if title:
                h = self.doc.add_heading(title, level=1)
                for run in h.runs:
                    run.font.name = self.template.font

    # ------------------------------------------------------------------
    # Generic primitives
    # ------------------------------------------------------------------

    def add_image(
        self,
        data_uri: str,
        *,
        title: str | None = None,
        caption: str | None = None,
        width_mm: float | None = None,
    ) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Mm

        if title:
            self.add_heading(title, level=3)

        img_bytes = decode_data_uri(data_uri)
        buf = io.BytesIO(img_bytes)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        width = Mm(width_mm) if width_mm else Inches(self.template.image_width_inches)
        run.add_picture(buf, width=width)

        if caption:
            self.add_caption(caption)

    def add_heading(self, text: str, level: int = 1) -> None:
        h = self.doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = self.template.font

    def add_paragraph(self, text: str) -> None:
        from docx.shared import Pt

        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self.template.font
        run.font.size = Pt(self.template.font_size_pt)

    def add_caption(self, text: str) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = self.template.font
        run.font.size = Pt(self.template.font_size_pt - 1)
        run.font.italic = True

    def add_table(self, headers: list[str], rows: list[list[str]]) -> None:
        n_cols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=n_cols)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        for r, row in enumerate(rows, start=1):
            for c, cell_value in enumerate(row):
                table.rows[r].cells[c].text = str(cell_value)

    def add_page_break(self) -> None:
        self.doc.add_page_break()

    # ------------------------------------------------------------------
    # Native escape hatch
    # ------------------------------------------------------------------

    def modify(self, fn: Callable[[Document], None]) -> None:
        """Run *fn* against the underlying python-docx ``Document``.

        Lets elements use python-docx features that aren't exposed via
        the generic primitives (custom styles, tables-of-contents,
        section breaks, headers/footers, etc.).
        """
        fn(self.doc)

    # ------------------------------------------------------------------
    # Finalize
    # ------------------------------------------------------------------

    def build(self) -> bytes:
        out = io.BytesIO()
        self.doc.save(out)
        return out.getvalue()
